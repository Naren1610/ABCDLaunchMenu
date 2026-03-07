import sys
from pypdf import PdfReader
from docx import Document

def extract_pdf(filepath):
    try:
        reader = PdfReader(filepath)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        return str(e)

def extract_docx(filepath):
    try:
        doc = Document(filepath)
        text = ""
        for p in doc.paragraphs:
            if p.text.strip():
                text += p.text + "\n"
        # also check tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.replace('\n', ' '))
                text += " | ".join(row_text) + "\n"
        return text
    except Exception as e:
        return str(e)

if __name__ == "__main__":
    with open("extracted_text.txt", "w", encoding="utf-8") as f:
        f.write("--- PDF CONTENT ---\n")
        f.write(extract_pdf("Launch Menu.pdf"))
        f.write("\n\n--- DOCX CONTENT ---\n")
        f.write(extract_docx("Launch Menu Prices.docx"))
